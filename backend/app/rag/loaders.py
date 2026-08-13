"""文档加载器工厂:PDF / DOCX / XLSX / TXT / MD → 统一 TextBlock 列表。

TextBlock = {text, page, section}:text 为原始文本,page 为页码(可空),section 为章节名(可空)。
"""
# 为什么统一成 TextBlock:切分器只认一种输入结构,格式差异全部在加载层消化
# 扩展新格式(如 PPT)只需注册一个加载函数,主流程零改动
# 注意:load_document 是同步 IO/CPU 密集函数,调用方(入库任务)应通过 to_thread 调用
import logging
from dataclasses import dataclass
from pathlib import Path

import charset_normalizer

from app.core.exceptions import BadRequestError

logger = logging.getLogger(__name__)


@dataclass
class TextBlock:
    # 统一块结构:section 在切分阶段作为 chunk 元数据入库,检索时展示"来自哪一章"
    # text 保持原始文本:切分规则由切分器决定,加载器不做二次处理
    text: str
    page: int | None = None
    section: str | None = None


# 加载文档:path=文件路径、file_type=扩展名(pdf/docx/xlsx/txt/md);返回非空文本块列表
def load_document(path: str | Path, file_type: str) -> list[TextBlock]:
    # 路径统一转 Path 对象:兼容字符串与 Path 两种调用方式
    path = Path(path)
    # 工厂模式:按扩展名分发到对应加载器,新增格式只需在此注册一行
    loaders = {
        "pdf": _load_pdf,
        "docx": _load_docx,
        "xlsx": _load_xlsx,
        "txt": _load_text,
        "md": _load_text,
    }
    loader = loaders.get(file_type)
    if loader is None:
        # 未注册类型直接报错:避免静默返回空列表,让入库任务卡在"0 块"的假成功
        raise BadRequestError(f"不支持的文件类型: {file_type}")
    blocks = loader(path)
    # 过滤全空白块:纯空白段落切不出有效 chunk,还白花 embedding 费用
    blocks = [b for b in blocks if b.text.strip()]
    # 块顺序即文档阅读顺序:后续 chunk_index 依此分配
    logger.info("加载 %s 完成: %d 个文本块", path.name, len(blocks))
    return blocks


# ---------- PDF ----------

# PDF 加载:PyMuPDF 逐页提取文本,页码从 1 开始与文档页码一致
def _load_pdf(path: Path) -> list[TextBlock]:
    import fitz  # PyMuPDF

    blocks: list[TextBlock] = []
    # with 保证文档句柄关闭,防止大文档场景句柄泄漏
    with fitz.open(path) as doc:
        # get_text("text") 按阅读顺序提取纯文本,带页码进 TextBlock
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            # 空页跳过:扫描件中的空白页不会产生文本块
            if text:
                # PDF 无结构化章节信息,section 恒为空,由切分阶段统一处理
                blocks.append(TextBlock(text=text, page=page_index))
    return blocks


# ---------- DOCX ----------

# DOCX 加载:段落与表格两路提取;样式名以 heading 开头的段落视为章节标题
def _load_docx(path: Path) -> list[TextBlock]:
    import docx

    d = docx.Document(path)
    blocks: list[TextBlock] = []
    # 当前章节标题:标题行之后的所有段落都归入该章节,便于定位"哪一章的内容"
    current_section: str | None = None

    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 样式名因文档而异(Heading 1/标题 1/heading1):统一转小写再做前缀匹配
        if para.style.name.lower().startswith("heading"):
            current_section = text
            continue
        blocks.append(TextBlock(text=text, section=current_section))

    # 表格:每行 → "列名:值" 形式,便于检索规格参数
    for table in d.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            # 空表跳过:没有内容可提取
            continue
        header = rows[0]
        for row in rows[1:]:
            # 键值对拼接:查询"型号 X100"可直接命中该行;zip 自动截断行单元格比表头短的情况
            # 表格块在段落之后追加:块顺序与文档阅读顺序保持一致
            line = "; ".join(f"{h}: {v}" for h, v in zip(header, row) if h and v)
            if line:
                # 无章节上下文时标记为"表格",检索结果仍能看出出处
                blocks.append(TextBlock(text=line, section=current_section or "表格"))
    return blocks


# ---------- XLSX ----------

# XLSX 加载:openpyxl 只读模式逐 sheet 处理,规格参数类表格是电商知识库的主力来源
def _load_xlsx(path: Path) -> list[TextBlock]:
    from openpyxl import load_workbook

    # read_only 省内存(大表格友好);data_only 取公式计算结果而非公式本身
    wb = load_workbook(path, read_only=True, data_only=True)
    blocks: list[TextBlock] = []
    for sheet in wb.worksheets:
        rows = sheet.iter_rows(values_only=True)
        header = None
        for row in rows:
            # 单元格统一 str() 化:数值(如价格 599)转成文本才可被检索匹配
            values = [str(c).strip() if c is not None else "" for c in row]
            if not any(values):
                # 跳过全空行:不产生无意义块
                continue
            if header is None:
                header = values  # 第一行作为表头,与用户对表格的直觉一致
                continue
            line = "; ".join(f"{h}: {v}" for h, v in zip(header, values) if h and v)
            if line:
                # section 用工作表名:检索结果能看出"来自哪个工作表"
                blocks.append(TextBlock(text=line, section=sheet.title))
    # read_only 模式下手动关闭:避免文件句柄滞留
    wb.close()
    return blocks


# ---------- TXT / MD ----------

def _decode_bytes(raw: bytes) -> str:
    """显式解码顺序:UTF-8 → GB18030(GBK 超集)→ 探测兜底。

    charset-normalizer 对短文本的 GBK 字节误判率较高,因此先尝试严格解码。
    """
    # 依次尝试严格解码:utf-8 优先(现代文件绝对主流,一次成功概率最高),gb18030 覆盖中文遗留编码
    for enc in ("utf-8", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    # 严格解码全失败才交给探测库兜底:它内部有统计启发式,更适合疑难编码
    detected = charset_normalizer.from_bytes(raw).best()
    if detected is not None:
        return str(detected)
    # 最后手段:替换非法字节为占位符,乱码入库总比中断入库流程好
    return raw.decode("utf-8", errors="replace")


# TXT/MD 加载:读字节 → 解码 → 按空行粗分段落;page/section 留空
# 同一函数服务 txt 与 md:两者解码策略一致,差异只在切分阶段体现
def _load_text(path: Path) -> list[TextBlock]:
    raw = path.read_bytes()
    text = _decode_bytes(raw)
    # 按空行分段:md/txt 的段落边界清晰,粗分交给切分器做更细的语义重组
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
    return [TextBlock(text=p) for p in paragraphs]
