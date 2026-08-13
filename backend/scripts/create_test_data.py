"""生成中文电商商品知识库测试文档(5 种格式,10 个事实锚点)。

产物:
    docs/data/samples/商品说明.md
    docs/data/samples/商品参数.txt     (故意用 GBK 编码,验证编码探测)
    docs/data/samples/商品清单.xlsx    (多行商品 + 规格列)
    docs/data/samples/商品手册.docx    (段落 + 表格)
    docs/data/samples/商品说明书.pdf   (中文 PDF,PyMuPDF 抽取)

事实锚点(可用作端到端验证的"标准答案"):
    1. 星云手机支持 7 天无理由退换货
    2. 星云手机电池容量 5000mAh
    3. 星云手机支持 66W 有线快充
    4. 星云手机 8GB+256GB 版本售价 2999 元
    5. 星云手机后置摄像头 5000 万像素
    6. 星云手机屏幕尺寸 6.7 英寸
    7. 星云手机支持 IP68 防水防尘
    8. 清风空气净化器适用于 40 平米房间
    9. 清风空气净化器噪音 33 分贝(睡眠模式)
    10. 星云手机购买后一年内非人为损坏免费维修
"""
# ---------------------------------------------------------------------------
# 测试数据生成器:产出 docs/data/samples 下 5 种格式的商品文档,
# 供 e2e_smoke.py 冒烟测试与本地手工验证使用。
# 内容围绕两款虚拟商品(星云手机 X1 / 清风空气净化器 K3)编写,
# 并埋入 10 个"事实锚点"——即问答时的标准答案,
# 测试断言据此判断检索是否命中、生成是否答对。
# ---------------------------------------------------------------------------
import sys
from pathlib import Path

# 把 backend 根目录临时加入模块搜索路径:
# 脚本由 `python -m scripts.create_test_data` 方式运行,
# 需要能 import 到 backend 根目录下的 app 包(如 app.rag.loaders)。
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# 输出目录与 e2e_smoke.py 里的 SAMPLES 常量指向同一位置,
# 两个脚本约定一致,冒烟测试才能找到这里生成的样本文件。
OUT_DIR = Path(__file__).resolve().parent.parent.parent / "docs" / "data" / "samples"

# ---- 商品说明.md(自然语言段落)----
# markdown 内容按"## 小标题"组织成 6 个主题段,
# 与文档加载器的 markdown 切分规则(标题作为 section)配合。
# 覆盖产品概述、退换货、电池、拍照、防护、保修等常见商品话题,
# 其中包含事实锚点 1~7 与 10(退换货天数、电池容量、快充功率、售价、像素、屏幕、防水、保修)。
MD_CONTENT = """# 星云手机 X1 商品说明

## 一、产品概述

星云手机 X1 是星云科技推出的旗舰智能手机,搭载高性能八核处理器,配备 6.7 英寸 AMOLED 全面屏,支持 120Hz 高刷新率,显示效果细腻流畅。

## 二、退换货政策

本商品支持 7 天无理由退换货。自签收之日起 7 天内,商品未拆封使用、配件齐全、不影响二次销售的情况下,可申请无理由退换货。质量问题导致的退换货不受 7 天限制。

## 三、电池与充电

星云手机 X1 内置 5000mAh 大容量电池,支持 66W 有线快充,30 分钟可充至 80% 电量,还支持 15W 无线充电和 10W 反向充电。重度使用场景下续航可达一整天。

## 四、拍照功能

后置摄像头为 5000 万像素主摄,搭配 1300 万像素超广角镜头和 200 万像素微距镜头;前置摄像头 3200 万像素,支持 AI 美颜和夜景模式。

## 五、防护等级

支持 IP68 防水防尘等级,可在一米水深下浸泡 30 分钟不受损坏,日常泼溅、雨淋无需担心。

## 六、购买与保修

8GB+256GB 版本售价 2999 元,12GB+512GB 版本售价 3999 元。购买后一年内非人为损坏可享受免费维修服务,支持全国联保。
"""

# ---- 商品参数.txt(键值对格式)----
# 与 md 的段落式排版形成对照,验证"键:值"行文本的切分与检索;
# 含事实锚点 8、9(适用面积 40 平米、噪音 33 分贝)。
TXT_CONTENT = """清风空气净化器 K3 规格参数
适用面积:40 平方米
颗粒物 CADR:400 立方米/小时
甲醛 CADR:150 立方米/小时
运行噪音:33 分贝(睡眠模式),48 分贝(标准模式)
滤芯寿命:约 6 个月
额定功率:45W
产品尺寸:320mm x 320mm x 620mm
产品重量:7.8kg
质保政策:整机一年质保,滤芯耗材不在保修范围内
使用提示:建议每 6 个月更换一次滤芯,以保证净化效果
"""

# ---- 商品清单.xlsx(表格格式)----
# 首行为表头,其余为数据行;数值(售价/库存)与文本混用,
# 验证 openpyxl 读取后单元格类型还原是否准确。
XLSX_ROWS = [
    ["商品名称", "型号", "颜色", "售价(元)", "库存", "发货方式"],
    ["星云手机 X1", "X1-8-256", "曜石黑", 2999, 500, "48小时内发货"],
    ["星云手机 X1", "X1-12-512", "星光银", 3999, 300, "48小时内发货"],
    ["清风空气净化器 K3", "K3-白", "珍珠白", 1299, 200, "现货速发"],
    ["清风空气净化器 K3", "K3-灰", "深空灰", 1299, 150, "现货速发"],
    ["星云蓝牙耳机 Pro", "Earpod-Pro", "经典黑", 399, 1000, "现货速发"],
]


# 以 UTF-8 无 BOM 写入:若带 BOM,加载器按 UTF-8 解码时首行会混入 ﻿ 字符,
# 导致检索时第一个事实锚点匹配不上。
def write_md() -> None:
    (OUT_DIR / "商品说明.md").write_text(MD_CONTENT, encoding="utf-8")


# 故意以 GBK 编码写入(Windows 记事本默认保存编码):
# 验证 txt 加载器依赖 charset-normalizer 自动探测编码,
# 而非写死 UTF-8,否则这类历史遗留文件会整篇乱码。
def write_txt_gbk() -> None:
    """故意用 GBK 编码写入,验证 charset-normalizer 编码探测。"""
    (OUT_DIR / "商品参数.txt").write_bytes(TXT_CONTENT.encode("gbk"))


# 使用 openpyxl 生成 .xlsx(Excel 2007+ 的 XML 格式):
# openpyxl 只支持 xlsx 一种格式,不支持旧版 .xls,
# 若项目需要兼容 .xls 需另选 xlwt/xlrd 等库。
def write_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    # 新工作簿默认自带一个空 sheet,直接改名复用,避免多出一个空工作表。
    ws.title = "商品清单"
    for row in XLSX_ROWS:
        ws.append(row)
    wb.save(OUT_DIR / "商品清单.xlsx")


# 用 python-docx 生成 .docx,包含标题、段落、表格三种元素:
# 验证 docx 加载器既能提取正文段落,也能把表格还原成结构化文本。
def write_docx() -> None:
    import docx

    d = docx.Document()
    d.add_heading("星云蓝牙耳机 Pro 产品手册", level=0)
    d.add_heading("产品特性", level=1)
    d.add_paragraph("星云蓝牙耳机 Pro 采用蓝牙 5.3 芯片,支持主动降噪(降噪深度 42dB),单次续航 8 小时,搭配充电仓总续航 36 小时。支持双设备连接,可在手机与平板之间无缝切换。")
    d.add_heading("规格参数", level=1)
    table = d.add_table(rows=5, cols=2)
    data = [
        ("蓝牙版本", "5.3"),
        ("降噪深度", "42dB"),
        ("单次续航", "8 小时"),
        ("充电仓总续航", "36 小时"),
        ("防水等级", "IPX5"),
    ]
    for i, (k, v) in enumerate(data):
        # 表格按"属性-值"两列逐行填充:docx 没有按行列名写值的便捷 API,
        # 只能先取到 cell 再逐个写 text。
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    d.add_heading("保修政策", level=1)
    d.add_paragraph("耳机主体保修一年,充电线及耳帽属于易耗品不在保修范围内。")
    d.save(OUT_DIR / "商品手册.docx")


# 用 reportlab 手工绘制 PDF 而非复用现有模板:
# 这样生成的 PDF 自带文本层,可验证 PyMuPDF 对中文文本的直接抽取,
# 而不是依赖 OCR 兜底(扫描版 PDF 才需要 OCR 识别)。
def write_pdf() -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    # reportlab 内置字体仅覆盖拉丁字符集,写中文必须先注册 TTF 中文字体,
    # 否则 drawString 遇到中文会抛错或渲染成空白方块。
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if font_path.exists():
        pdfmetrics.registerFont(TTFont("MSYH", str(font_path)))
    else:
        # 微软雅黑不存在(如 Linux/CI 环境)时回退到黑体,保证脚本跨平台可用。
        pdfmetrics.registerFont(TTFont("MSYH", "C:/Windows/Fonts/simhei.ttf"))

    c = canvas.Canvas(str(OUT_DIR / "商品说明书.pdf"), pagesize=A4)
    c.setFont("MSYH", 14)
    lines = [
        "星云手机 X1 使用说明书",
        "",
        "一、安全须知",
        "请勿将手机长时间放置在高温环境下,避免阳光直射。充电时请使用原装 66W 充电器。",
        "",
        "二、电池使用",
        "本机内置 5000mAh 电池,支持 66W 有线快充。建议使用原装充电器以保证充电速度和安全性。",
        "",
        "三、防水说明",
        "手机支持 IP68 防水防尘,但请勿在带水状态下插拔充电线。",
        "",
        "四、保修说明",
        "自购买之日起一年内,非人为损坏可享受免费维修。",
    ]
    y = 780
    for line in lines:
        c.drawString(72, y, line)
        # PDF 画布没有流式排版,必须手工维护基线坐标,每次换行递减一个行高。
        y -= 28
    c.showPage()
    c.save()


# 先确保目录存在:首次运行或目录被清理后,直接写文件会抛 FileNotFoundError。
def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_md()
    write_txt_gbk()
    write_xlsx()
    write_docx()
    write_pdf()
    print(f"测试数据已生成到: {OUT_DIR}")
    print("共 5 个文件,内含 10 个事实锚点(见脚本头部注释)")


if __name__ == "__main__":
    # 模块级守卫:只有被直接执行时才跑 main,
    # 被其它脚本 import 时不产生副作用(例如复用 MD_CONTENT 常量做断言)。
    main()
